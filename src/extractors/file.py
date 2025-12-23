"""
File data extractor for SuperInsight Platform.

Provides extraction from PDF, Word, HTML, and text files.
"""

import logging
import os
from typing import List, Optional, Dict, Any
from datetime import datetime
from pathlib import Path
import mimetypes

# Optional file processing libraries - import only if available
try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

import requests
from urllib.parse import urlparse, urljoin

from src.extractors.base import BaseExtractor, FileConfig, ExtractionResult, FileType
from src.models.document import Document

logger = logging.getLogger(__name__)


class FileExtractor(BaseExtractor):
    """File extractor for various document formats."""
    
    def __init__(self, config: FileConfig):
        super().__init__(config)
        self.config: FileConfig = config
    
    def test_connection(self) -> bool:
        """Test file accessibility."""
        try:
            if self.config.file_path.startswith(('http://', 'https://')):
                # Test web URL accessibility
                response = requests.head(
                    self.config.file_path,
                    timeout=self.config.connection_timeout,
                    verify=self.config.verify_ssl
                )
                return response.status_code == 200
            else:
                # Test local file accessibility
                return os.path.exists(self.config.file_path) and os.path.isfile(self.config.file_path)
        except Exception as e:
            logger.error(f"File connection test failed: {e}")
            return False
    
    def extract_data(self, **kwargs) -> ExtractionResult:
        """Extract data from file based on file type."""
        try:
            if self.config.file_type == FileType.PDF:
                return self._extract_pdf()
            elif self.config.file_type == FileType.DOCX:
                return self._extract_docx()
            elif self.config.file_type == FileType.TXT:
                return self._extract_text()
            elif self.config.file_type == FileType.HTML:
                return self._extract_html()
            else:
                return ExtractionResult(
                    success=False,
                    error=f"Unsupported file type: {self.config.file_type}"
                )
        except Exception as e:
            logger.error(f"File extraction failed: {e}")
            return ExtractionResult(success=False, error=str(e))
    
    def _extract_pdf(self) -> ExtractionResult:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            return ExtractionResult(
                success=False,
                error="pypdf is required for PDF extraction. Install with: pip install pypdf"
            )
        
        try:
            documents = []
            
            if self.config.file_path.startswith(('http://', 'https://')):
                # Download PDF from URL
                response = requests.get(
                    self.config.file_path,
                    timeout=self.config.read_timeout,
                    verify=self.config.verify_ssl
                )
                response.raise_for_status()
                
                # Create temporary file-like object
                from io import BytesIO
                pdf_file = BytesIO(response.content)
            else:
                # Read local PDF file
                pdf_file = open(self.config.file_path, 'rb')
            
            try:
                pdf_reader = pypdf.PdfReader(pdf_file)
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    
                    if text.strip():  # Only create document if page has text
                        document = Document(
                            source_type="file",
                            source_config={
                                "file_path": self.config.file_path,
                                "file_type": "pdf",
                                "page_number": page_num + 1
                            },
                            content=text,
                            metadata={
                                "file_name": os.path.basename(self.config.file_path),
                                "page_number": page_num + 1,
                                "total_pages": len(pdf_reader.pages),
                                "extraction_time": datetime.now().isoformat()
                            }
                        )
                        documents.append(document)
                
                logger.info(f"Extracted {len(documents)} pages from PDF")
                return ExtractionResult(success=True, documents=documents)
                
            finally:
                if hasattr(pdf_file, 'close'):
                    pdf_file.close()
                    
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ExtractionResult(success=False, error=str(e))
    
    def _extract_docx(self) -> ExtractionResult:
        """Extract text from Word document."""
        if not DOCX_AVAILABLE:
            return ExtractionResult(
                success=False,
                error="python-docx is required for DOCX extraction. Install with: pip install python-docx"
            )
        
        try:
            if self.config.file_path.startswith(('http://', 'https://')):
                # Download DOCX from URL
                response = requests.get(
                    self.config.file_path,
                    timeout=self.config.read_timeout,
                    verify=self.config.verify_ssl
                )
                response.raise_for_status()
                
                from io import BytesIO
                docx_file = BytesIO(response.content)
                doc = DocxDocument(docx_file)
            else:
                # Read local DOCX file
                doc = DocxDocument(self.config.file_path)
            
            # Extract paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Combine all paragraphs into single document
            content = '\n'.join(paragraphs)
            
            if content.strip():
                document = Document(
                    source_type="file",
                    source_config={
                        "file_path": self.config.file_path,
                        "file_type": "docx"
                    },
                    content=content,
                    metadata={
                        "file_name": os.path.basename(self.config.file_path),
                        "paragraph_count": len(paragraphs),
                        "extraction_time": datetime.now().isoformat()
                    }
                )
                
                logger.info(f"Extracted DOCX with {len(paragraphs)} paragraphs")
                return ExtractionResult(success=True, documents=[document])
            else:
                return ExtractionResult(success=False, error="No text content found in DOCX")
                
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ExtractionResult(success=False, error=str(e))
    
    def _extract_text(self) -> ExtractionResult:
        """Extract content from plain text file."""
        try:
            if self.config.file_path.startswith(('http://', 'https://')):
                # Download text from URL
                response = requests.get(
                    self.config.file_path,
                    timeout=self.config.read_timeout,
                    verify=self.config.verify_ssl
                )
                response.raise_for_status()
                content = response.text
            else:
                # Read local text file
                with open(self.config.file_path, 'r', encoding=self.config.encoding) as f:
                    content = f.read()
            
            if content.strip():
                document = Document(
                    source_type="file",
                    source_config={
                        "file_path": self.config.file_path,
                        "file_type": "txt",
                        "encoding": self.config.encoding
                    },
                    content=content,
                    metadata={
                        "file_name": os.path.basename(self.config.file_path),
                        "file_size": len(content),
                        "extraction_time": datetime.now().isoformat()
                    }
                )
                
                logger.info(f"Extracted text file with {len(content)} characters")
                return ExtractionResult(success=True, documents=[document])
            else:
                return ExtractionResult(success=False, error="No content found in text file")
                
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return ExtractionResult(success=False, error=str(e))
    
    def _extract_html(self) -> ExtractionResult:
        """Extract text from HTML content."""
        if not HTML_AVAILABLE:
            return ExtractionResult(
                success=False,
                error="beautifulsoup4 is required for HTML extraction. Install with: pip install beautifulsoup4"
            )
        
        try:
            if self.config.file_path.startswith(('http://', 'https://')):
                # Download HTML from URL
                response = requests.get(
                    self.config.file_path,
                    timeout=self.config.read_timeout,
                    verify=self.config.verify_ssl,
                    headers={'User-Agent': 'SuperInsight-Platform/1.0'}
                )
                response.raise_for_status()
                html_content = response.text
            else:
                # Read local HTML file
                with open(self.config.file_path, 'r', encoding=self.config.encoding) as f:
                    html_content = f.read()
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            if text.strip():
                # Extract additional metadata
                title = soup.find('title')
                title_text = title.get_text() if title else ""
                
                meta_description = soup.find('meta', attrs={'name': 'description'})
                description = meta_description.get('content', '') if meta_description else ""
                
                document = Document(
                    source_type="file",
                    source_config={
                        "file_path": self.config.file_path,
                        "file_type": "html"
                    },
                    content=text,
                    metadata={
                        "file_name": os.path.basename(self.config.file_path),
                        "title": title_text,
                        "description": description,
                        "content_length": len(text),
                        "extraction_time": datetime.now().isoformat()
                    }
                )
                
                logger.info(f"Extracted HTML with {len(text)} characters")
                return ExtractionResult(success=True, documents=[document])
            else:
                return ExtractionResult(success=False, error="No text content found in HTML")
                
        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            return ExtractionResult(success=False, error=str(e))


class WebExtractor(BaseExtractor):
    """Web content extractor for crawling web pages."""
    
    def __init__(self, base_url: str, max_pages: int = 10):
        from src.extractors.base import APIConfig
        config = APIConfig(
            base_url=base_url,
            headers={'User-Agent': 'SuperInsight-Platform/1.0'}
        )
        super().__init__(config)
        self.base_url = base_url
        self.max_pages = max_pages
        self.visited_urls = set()
    
    def test_connection(self) -> bool:
        """Test web site accessibility."""
        try:
            response = requests.head(
                self.base_url,
                timeout=self.config.connection_timeout,
                verify=self.config.verify_ssl,
                headers=self.config.headers
            )
            return response.status_code == 200
        except Exception as e:
            logger.error(f"Web connection test failed: {e}")
            return False
    
    def extract_data(self, max_depth: int = 2, **kwargs) -> ExtractionResult:
        """Extract content from web pages with optional crawling."""
        try:
            documents = []
            urls_to_visit = [self.base_url]
            current_depth = 0
            
            while urls_to_visit and current_depth < max_depth and len(documents) < self.max_pages:
                current_urls = urls_to_visit.copy()
                urls_to_visit = []
                
                for url in current_urls:
                    if url in self.visited_urls or len(documents) >= self.max_pages:
                        continue
                    
                    result = self._extract_single_page(url)
                    if result.success and result.documents:
                        documents.extend(result.documents)
                        self.visited_urls.add(url)
                        
                        # Find links for next depth level
                        if current_depth < max_depth - 1:
                            links = self._extract_links(url)
                            urls_to_visit.extend(links)
                
                current_depth += 1
            
            logger.info(f"Extracted {len(documents)} web pages")
            return ExtractionResult(success=True, documents=documents)
            
        except Exception as e:
            logger.error(f"Web extraction failed: {e}")
            return ExtractionResult(success=False, error=str(e))
    
    def _extract_single_page(self, url: str) -> ExtractionResult:
        """Extract content from a single web page."""
        try:
            # Create file config for HTML extraction
            file_config = FileConfig(
                file_path=url,
                file_type=FileType.HTML,
                connection_timeout=self.config.connection_timeout,
                read_timeout=self.config.read_timeout,
                use_ssl=self.config.use_ssl,
                verify_ssl=self.config.verify_ssl
            )
            
            file_extractor = FileExtractor(file_config)
            return file_extractor.extract_data()
            
        except Exception as e:
            logger.error(f"Failed to extract page {url}: {e}")
            return ExtractionResult(success=False, error=str(e))
    
    def _extract_links(self, url: str) -> List[str]:
        """Extract links from a web page for crawling."""
        try:
            response = requests.get(
                url,
                timeout=self.config.read_timeout,
                verify=self.config.verify_ssl,
                headers=self.config.headers
            )
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            links = []
            
            for link in soup.find_all('a', href=True):
                href = link['href']
                
                # Convert relative URLs to absolute
                absolute_url = urljoin(url, href)
                
                # Only include links from the same domain
                if urlparse(absolute_url).netloc == urlparse(self.base_url).netloc:
                    links.append(absolute_url)
            
            return list(set(links))  # Remove duplicates
            
        except Exception as e:
            logger.error(f"Failed to extract links from {url}: {e}")
            return []


def detect_file_type(file_path: str) -> Optional[FileType]:
    """Detect file type from file path or URL."""
    try:
        # Get file extension
        if file_path.startswith(('http://', 'https://')):
            # For URLs, try to get extension from path
            parsed = urlparse(file_path)
            path = parsed.path
        else:
            path = file_path
        
        extension = Path(path).suffix.lower()
        
        # Map extensions to file types
        extension_map = {
            '.pdf': FileType.PDF,
            '.docx': FileType.DOCX,
            '.doc': FileType.DOCX,  # Treat .doc as .docx for now
            '.txt': FileType.TXT,
            '.html': FileType.HTML,
            '.htm': FileType.HTML,
        }
        
        return extension_map.get(extension)
        
    except Exception as e:
        logger.error(f"Failed to detect file type for {file_path}: {e}")
        return None